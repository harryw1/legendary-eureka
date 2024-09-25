import os
import threading
import tkinter as tk
import zipfile
from idlelib.tooltip import Hovertip
from tkinter import filedialog, messagebox, ttk
from xml.etree import ElementTree as ET

import docx
import openpyxl


def select_input_files():
    file_paths = filedialog.askopenfilenames(
        title="Select Word Documents",
        filetypes=[("Word Documents", "*.docx")]
    )
    if file_paths:
        input_files_list.delete(1.0, tk.END)
        input_files_list.insert(tk.END, '\n'.join(file_paths))

def select_output_file():
    file_path = filedialog.asksaveasfilename(
        defaultextension=".xlsx",
        filetypes=[("Excel Workbook", "*.xlsx")],
        title="Save Excel File As"
    )
    if file_path:
        output_file_var.set(file_path)

def clear_input_files():
    input_files_list.delete(1.0, tk.END)

def start_extraction():
    threading.Thread(target=extract_comments, daemon=True).start()

def extract_comments():
    input_files = input_files_list.get(1.0, tk.END).strip().split('\n')
    output_file = output_file_var.get()

    if not input_files or not output_file:
        messagebox.showerror("Error", "Please select input files and an output file.")
        return

    wb = openpyxl.Workbook()
    ws = wb.active
    if ws is None:
        ws = wb.create_sheet()
    ws.append(["File Name", "Comment Author", "Comment Text", "Page Number", "Referenced Text"])

    errors = []
    total_files = len(input_files)
    for i, file_path in enumerate(input_files):
        if not os.path.exists(file_path):
            errors.append(f"File not found: {file_path}")
            continue
        try:
            comments = extract_comments_from_docx(file_path)
            for comment in comments:
                ws.append([
                    os.path.basename(file_path),
                    comment['author'],
                    comment['text'],
                    comment['page'],
                    comment['referenced_text']
                ])
        except Exception as e:
            errors.append(f"Error processing file {file_path}:\n{str(e)}")
        progress_var.set((i + 1) / total_files * 100)
        root.update_idletasks()

    try:
        wb.save(output_file)
        messagebox.showinfo("Success", f"Comments extracted and saved to {output_file}")
    except Exception as e:
        errors.append(f"Error saving output file:\n{str(e)}")

    if errors:
        error_message = "\n\n".join(errors)
        messagebox.showerror("Errors Occurred", f"The following errors occurred:\n\n{error_message}")

def extract_comments_from_docx(file_path):
    comments = []
    doc = docx.Document(file_path)

    # Create a mapping of comment_id to comment details
    comment_map = {}
    with zipfile.ZipFile(file_path) as zip_file:  # Changed 'docx' to 'zip_file'
        try:
            comments_xml = zip_file.read("word/comments.xml")
            root = ET.fromstring(comments_xml)
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            for comment in root.findall('w:comment', ns):
                comment_id = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                author = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author')
                texts = []
                for p in comment.findall('w:p', ns):
                    for t in p.findall('.//w:t', ns):
                        texts.append(t.text if t.text else '')
                comment_text = ''.join(texts)
                comment_map[comment_id] = {'author': author, 'text': comment_text}
        except KeyError:
            # No comments in the document
            pass

    # Extract comments with page numbers and referenced text
    for i, paragraph in enumerate(doc.paragraphs):
        for run in paragraph.runs:
            comment_reference = run._element.find('.//w:commentReference', ns)
            if comment_reference is not None:
                comment_id = comment_reference.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                if comment_id in comment_map:
                    comment = comment_map[comment_id]
                    page_number = get_page_number(doc, i)
                    referenced_text = paragraph.text
                    comments.append({
                        'author': comment['author'],
                        'text': comment['text'],
                        'page': page_number,
                        'referenced_text': referenced_text
                    })

    return comments


def get_page_number(doc, paragraph_index):
    page_number = 1
    current_page_start = 0
    for i, p in enumerate(doc.paragraphs):
        if i == paragraph_index:
            return page_number
        if p.runs and p.runs[0].element.tag.endswith('br') and p.runs[0].element.get('type') == 'page':
            page_number += 1
            current_page_start = i + 1
    return page_number

# GUI setup with ttk
root = tk.Tk()
root.title("Word Comments Extractor")

# Styles for ttk widgets
style = ttk.Style()
style.configure("TLabel", font=("Helvetica", 12))
style.configure("TButton", font=("Helvetica", 12))
style.configure("TFrame", background="#f0f0f0")  # Light gray background for frames

# Main frame
mainframe = ttk.Frame(root, padding="10")
mainframe.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

# Input files frame
input_frame = ttk.LabelFrame(mainframe, text="Input Files", padding="5")
input_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=5)

input_files_scrollbar = ttk.Scrollbar(input_frame)
input_files_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

input_files_list = tk.Text(input_frame, height=10, width=50, yscrollcommand=input_files_scrollbar.set)
input_files_list.pack(expand=True, fill=tk.BOTH)
input_files_scrollbar.config(command=input_files_list.yview)

# Output file frame
output_frame = ttk.LabelFrame(mainframe, text="Output File", padding="5")
output_frame.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=5)

output_file_var = tk.StringVar()
ttk.Entry(output_frame, textvariable=output_file_var, width=50).pack(side=tk.LEFT, padx=5)
select_output_btn = ttk.Button(output_frame, text="Select File", command=select_output_file)
select_output_btn.pack(side=tk.LEFT, padx=5)

# Progress frame
progress_frame = ttk.LabelFrame(mainframe, text="Progress", padding="5")
progress_frame.grid(row=2, column=0, sticky=(tk.W, tk.E), pady=5)

progress_var = tk.DoubleVar()
progress_bar = ttk.Progressbar(progress_frame, variable=progress_var, maximum=100)
progress_bar.pack(fill=tk.X, padx=5)
progress_label = ttk.Label(progress_frame, text="Ready")
progress_label.pack()

# Buttons
button_frame = ttk.Frame(mainframe, padding="5")
button_frame.grid(row=3, column=0, pady=10)

select_files_btn = ttk.Button(button_frame, text="Select Files", command=select_input_files)
select_files_btn.pack(side=tk.LEFT, padx=5)

clear_files_btn = ttk.Button(button_frame, text="Clear Files", command=clear_input_files)
clear_files_btn.pack(side=tk.LEFT, padx=5)

extract_btn = ttk.Button(button_frame, text="Extract Comments", command=start_extraction)
extract_btn.pack(side=tk.LEFT, padx=5)

# Configure grid resizing
root.columnconfigure(0, weight=1)
root.rowconfigure(0, weight=1)
mainframe.columnconfigure(0, weight=1)
mainframe.rowconfigure(0, weight=1)
mainframe.rowconfigure(1, weight=0)
mainframe.rowconfigure(2, weight=0)
mainframe.rowconfigure(3, weight=0)

root.mainloop()
